import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generar_word_asis():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    salidas_dir = os.path.join(base_dir, 'salidas')
    out_docx = os.path.join(base_dir, 'informe_situacion_salud_tres_de_febrero.docx')
    
    print("--- CREANDO DOCUMENTO WORD CON FUENTES Y REFERENCIAS ACADÉMICAS EN CADA TABLA Y GRÁFICO ---")
    doc = Document()
    
    # Configurar márgenes de página
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
        
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # --- PORTADA OFICIAL ---
    p_portada_inst = doc.add_paragraph()
    p_portada_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_portada_inst.add_run("MUNICIPALIDAD DE TRES DE FEBRERO\nSECRETARÍA DE SALUD Y GESTIÓN OPERATIVA\n")
    run_inst.font.bold = True
    run_inst.font.size = Pt(12)
    run_inst.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    p_portada_tit = doc.add_paragraph()
    p_portada_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_portada_tit.paragraph_format.space_before = Pt(24)
    p_portada_tit.paragraph_format.space_after = Pt(12)
    run_tit = p_portada_tit.add_run("INFORME TÉCNICO Y ANÁLISIS DE SITUACIÓN DE SALUD (ASIS)\nDEL MUNICIPIO DE TRES DE FEBRERO")
    run_tit.font.bold = True
    run_tit.font.size = Pt(20)
    run_tit.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_portada_sub = doc.add_paragraph()
    p_portada_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_portada_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_portada_sub.add_run("Diagnóstico Cartográfico Georreferenciado, Infraestructura Sanitaria,\nEvolución Epidemiológica y Red de Efectores para Elevación a Región Sanitaria VII")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Tabla de metadatos de portada
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Entidad Emisora:", "Secretaría de Salud — Municipalidad de Tres de Febrero (Partido 06840)"),
        ("Destinatario:", "Dirección y Consejo Técnico-Epidemiológico — Región Sanitaria VII (Ministerio de Salud PBA)"),
        ("Fuentes de Datos:", "Censo 2022 INDEC (Definitivo), Cartografía Oficial INDEC/IGN/PBA, SISA, DEIS, UNGS (2012-2019)"),
        ("Fecha de Emisión:", "Julio 2026 — Documento Oficial de Gestión e Integración Regional")
    ]
    for i, (label, val) in enumerate(meta_data):
        row_cells = table_meta.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = val
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        set_cell_background(row_cells[0], "F8FAFC")
        set_cell_background(row_cells[1], "F8FAFC")
        set_cell_margins(row_cells[0], 120, 120, 150, 150)
        set_cell_margins(row_cells[1], 120, 120, 150, 150)
        
    doc.add_page_break()
    
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix + " ")
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(text)
        return p

    def add_fuente_nota(fuente_texto):
        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(2)
        p_f.paragraph_format.space_after = Pt(12)
        p_f.paragraph_format.line_spacing = 1.05
        rf = p_f.add_run(f"📌 Fuente y Referencia: {fuente_texto}")
        rf.font.size = Pt(9)
        rf.font.italic = True
        rf.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        return p_f

    def insert_img(filename, caption, fuente="", width_in=6.2):
        img_path = os.path.join(salidas_dir, filename)
        if not os.path.exists(img_path):
            img_path = os.path.join(base_dir, filename)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(2)
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=Inches(width_in))
            
            p_cap = doc.add_paragraph(caption)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(2)
            rcap = p_cap.runs[0]
            rcap.font.size = Pt(9.5)
            rcap.font.bold = True
            rcap.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
            if fuente:
                add_fuente_nota(fuente)

    # --- SECCIÓN 1: MARCO METODOLÓGICO Y CARACTERIZACIÓN GENERAL DEL TERRITORIO ---
    add_h1("1. Marco Metodológico y Caracterización General del Municipio")
    add_p("Tres de Febrero es un municipio del primer cordón del Gran Buenos Aires, lindante directamente con la Ciudad Autónoma de Buenos Aires al oeste de la capital del país. Forma parte de una de las zonas más dinámicas del área metropolitana, caracterizándose por una fuerte conectividad e infraestructura interurbana y una cercanía estratégica a los principales centros urbanos y productivos.", "Contexto Geográfico y Dinámica Metropolitana:")
    add_p("Nuestra ciudad cabecera es Caseros, donde se concentra el centro administrativo, político y urbano del distrito, así como una amplia red de servicios, instituciones sanitarias y espacios públicos de referencia comunitaria.", "Ciudad Cabecera:")
    add_p("El presente documento integra y unifica el relevamiento exhaustivo del equipo de salud municipal con los microdatos definitivos del Censo Nacional 2022 (INDEC) y las series históricas epidemiológicas del Ministerio de Salud provincial. Para una representatividad espacial total, se importó la capa oficial poligonal con los 457 radios censales del partido (`06840`).", "Metodología e Integración de Fuentes:")
    
    # Tabla Indicadores Básicos
    t_ind = doc.add_table(rows=4, cols=2)
    t_ind.alignment = WD_TABLE_ALIGNMENT.CENTER
    inds = [
        ("Población Total (Censo 2022 INDEC)", "364.176 habitantes (Crecimiento intercensal +7,09%)"),
        ("Superficie Distrital", "45 km² (Territorio enteramente urbano e industrial)"),
        ("Densidad Poblacional", "8.021,5 hab/km² — 4° municipio más denso de toda la República Argentina"),
        ("Estructura de Género", "191.214 mujeres (52,5%) y 172.962 varones (47,5%)")
    ]
    for idx, (lab, v) in enumerate(inds):
        t_ind.rows[idx].cells[0].text = lab
        t_ind.rows[idx].cells[1].text = v
        t_ind.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(t_ind.rows[idx].cells[0], "F8FAFC")
        set_cell_margins(t_ind.rows[idx].cells[0], 80, 80, 100, 100)
        set_cell_margins(t_ind.rows[idx].cells[1], 80, 80, 100, 100)
    add_fuente_nota("INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Resultados definitivos de población por sexos y densidad territorial del Partido de Tres de Febrero (`06840`) e Instituto Geográfico Nacional (IGN).")

    # --- SECCIÓN 2: EVOLUCIÓN EPIDEMIOLÓGICA E INDICADORES SANITARIOS HISTÓRICOS ---
    add_h1("2. Evolución Epidemiológica e Indicadores Sanitarios de la Última Década")
    add_p("A lo largo de los últimos diez años, el Municipio de Tres de Febrero evidencia una mejora sostenida y estadísticamente significativa en sus indicadores materno-infantiles y sanitarios primarios, fruto de la consolidación del primer nivel de atención y el control prenatal temprano en los CAPS.", "Logros Sanitarios Sustantivos:")
    add_p("Según las series estadísticas consolidadas por el Ministerio de Salud de la Provincia de Buenos Aires y el Instituto del Conurbano de la UNGS (serie 2012–2019):", "Indicadores de Impacto:")
    
    t_epi = doc.add_table(rows=4, cols=3)
    t_epi.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_epi = ["Indicador Epidemiológico", "Valor Inicial (2012)", "Valor Alcanzado / Actual"]
    for j, h in enumerate(headers_epi):
        t_epi.rows[0].cells[j].text = h
        t_epi.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_epi.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(t_epi.rows[0].cells[j], "1E3A8A")
        set_cell_margins(t_epi.rows[0].cells[j], 80, 80, 80, 80)
        
    epi_data = [
        ("Tasa de Mortalidad Infantil (TMI)", "9,2 por cada 1.000 nacidos vivos", "7,1 por cada 1.000 nacidos vivos (Descenso notable)"),
        ("Porcentaje de Maternidad Adolescente", "10,1% sobre el total de nacimientos", "7,2% sobre el total de nacimientos (Reducción continua)"),
        ("Tasa de Mortalidad Materna", "Baja incidencia distrital", "Prácticamente nula (0,0 por cada 10.000 nacidos vivos)")
    ]
    for row_idx, row_val in enumerate(epi_data):
        for col_idx, text_val in enumerate(row_val):
            t_epi.rows[row_idx+1].cells[col_idx].text = text_val
            set_cell_background(t_epi.rows[row_idx+1].cells[col_idx], "F8FAFC")
            set_cell_margins(t_epi.rows[row_idx+1].cells[col_idx], 70, 70, 80, 80)
            if col_idx == 2:
                t_epi.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.bold = True
                t_epi.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    add_fuente_nota("Dirección de Estadísticas e Información en Salud (DEIS - Ministerio de Salud de la Nación), Región Sanitaria VII (Ministerio de Salud PBA) e Instituto del Conurbano (ICO-UNGS, Series 2012–2019).")

    # --- SECCIÓN 3: ESTRUCTURA DEMOGRÁFICA, ENVEJECIMIENTO Y PREVISIÓN SOCIAL ---
    add_h1("3. Estructura Demográfica, Envejecimiento y Previsión Social (Censo 2022)")
    add_p("El rasgo demográfico y epidemiológico más concluyente de Tres de Febrero es su avanzado grado de envejecimiento poblacional. Sobre el universo de 362.319 personas en viviendas particulares censadas por el INDEC, el 20,77% (75.242 personas) percibe jubilación o pensión, proporción significativamente superior al promedio del Gran Buenos Aires (17,5%) y de la Provincia de Buenos Aires (18,3%).", "Transición Gerontológica:")
    
    insert_img('grafico_piramide_demografica_3f.png', 
               'Gráfico 1: Estructura Poblacional por Sexo, Grandes Franjas Etarias e Impacto Gerontológico en Tres de Febrero.',
               fuente="INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Estructura de población censada en viviendas particulares por sexo registrado y grandes grupos de edad en Tres de Febrero (`06840`).")
    
    add_h2("Desglose Previsional — Percepción de Jubilación y Pensión (Censo 2022)")
    t_prev = doc.add_table(rows=7, cols=3)
    t_prev.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_prev = ["Categoría Previsional (Censo 2022 INDEC)", "Personas", "Porcentaje sobre Población"]
    for j, h in enumerate(headers_prev):
        t_prev.rows[0].cells[j].text = h
        t_prev.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_prev.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(t_prev.rows[0].cells[j], "1E3A8A")
        set_cell_margins(t_prev.rows[0].cells[j], 80, 80, 80, 80)
        
    prev_data = [
        ("Total que Percibe Jubilación o Pensión", "75.242", "20,77% (Alta carga de ECNT)"),
        (" — Solo Jubilación", "47.171", "13,02%"),
        (" — Solo Pensión por fallecimiento", "4.546", "1,25%"),
        (" — Jubilación y Pensión por fallecimiento simultáneamente", "13.369", "3,69%"),
        (" — Solo Pensión de otro tipo (No contributiva / Discapacidad)", "10.156", "2,80%"),
        ("No Percibe Jubilación ni Pensión (< 65 años / Activos)", "287.077", "79,23%")
    ]
    for row_idx, (c1, c2, c3) in enumerate(prev_data):
        cells = t_prev.rows[row_idx+1].cells
        cells[0].text = c1
        cells[1].text = c2
        cells[2].text = c3
        if row_idx == 0:
            for c in cells:
                c.paragraphs[0].runs[0].font.bold = True
                set_cell_background(c, "FEF3C7")
        else:
            for c in cells:
                set_cell_background(c, "F8FAFC")
        for c in cells:
            set_cell_margins(c, 60, 60, 80, 80)
    add_fuente_nota("INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Población en viviendas particulares según condición de percepción de jubilación y/o pensión en el Partido de Tres de Febrero.")

    doc.add_page_break()

    # --- SECCIÓN 4: COBERTURA SANITARIA Y VULNERABILIDAD TERRITORIAL ---
    add_h1("4. Cobertura Sanitaria y Vulnerabilidad Desagregada por Localidad")
    add_p("El 72,8% de la población distrital cuenta con alguna cobertura de salud formal (70,76% con obra social o prepaga y 2,04% con programas o planes estatales). La proporción de población sin cobertura formal en Tres de Febrero (27,20% — 98.540 personas) es netamente inferior al promedio de los 24 Partidos del Gran Buenos Aires (37,60%) y al total de la Provincia (35,10%), posicionando al distrito favorablemente en el escenario general del conurbano.", "Panorama General de Aseguramiento:")
    
    t_cob = doc.add_table(rows=4, cols=3)
    t_cob.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Jurisdicción / Ámbito Territorial", "Población Sin Cobertura Formal", "Porcentaje de Dependencia Pública"]):
        t_cob.rows[0].cells[j].text = h
        t_cob.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_cob.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(t_cob.rows[0].cells[j], "1E3A8A")
        set_cell_margins(t_cob.rows[0].cells[j], 80, 80, 80, 80)
        
    cob_comp = [
        ("Municipio de Tres de Febrero", "98.540 personas", "27,2% (Favorable vs GBA)"),
        ("Promedio 24 Partidos del Gran Buenos Aires", "4.056.187 personas", "37,6%"),
        ("Total Provincia de Buenos Aires", "6.111.393 personas", "35,1%")
    ]
    for r_idx, (c1, c2, c3) in enumerate(cob_comp):
        cells = t_cob.rows[r_idx+1].cells
        cells[0].text = c1; cells[1].text = c2; cells[2].text = c3
        if r_idx == 0:
            cells[2].paragraphs[0].runs[0].font.bold = True
            cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
        for c in cells: set_cell_margins(c, 70, 70, 80, 80); set_cell_background(c, "F8FAFC")
    add_fuente_nota("INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Población en viviendas particulares según tipo de cobertura sanitaria por partido y provincia.")

    add_p("No obstante el promedio favorable distrital, el análisis espacial intramunicipal expone una profunda brecha de vulnerabilidad en el corredor Norte. Las localidades de El Libertador, Churruca, Pablo Podestá, Remedios de Escalada y Loma Hermosa concentran tasas de dependencia exclusiva del sistema público que oscilan entre el 38% y el 43%, demandando la focalización asistencial prioritaria del Hospital Bocalandro y los CAPS 10 a 13.", "Gradiente Territorial Norte-Sur:")
    
    insert_img('grafico_dependencia_por_localidad_3f.png', 
               'Gráfico 2: Gradiente Territorial de Dependencia del Sistema Público por las 15 Localidades vs Promedios GBA/PBA.',
               fuente="INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Procesamiento espacial por radios censales agrupados según las 15 localidades del Partido de Tres de Febrero en comparación con los 24 partidos del GBA.")

    # --- SECCIÓN 5: CARTOGRAFÍA POLIGONAL REAL CONTIGUA (CERO HUECOS) ---
    add_h1("5. Cartografía Poligonal Georreferenciada (457 Radios Censales)")
    add_p("El mapa de coropletas se construyó importando e integrando la totalidad de los 457 radios censales (`RADIO_CENS`) correspondientes al código INDEC `06840` sobre la cartografía oficial de la Provincia de Buenos Aires.", "Geometría Territorial Exacta:")
    add_p("El resultado cartográfico garantiza la representación de la malla urbana distrital y muestra la ubicación georreferenciada de la red de 19 efectores sanitarios provinciales y municipales.", "Precisión Territorial SIG:")
    
    insert_img('mapa_sistema_publico_3f.png', 
               'Mapa 1: Cartografía Georreferenciada de los 457 Radios Censales de Tres de Febrero y Red de 19 Efectores sanitarios.',
               fuente="Cartografía Poligonal Oficial del Instituto Nacional de Estadística y Censos (INDEC) e Instituto Geográfico Nacional (IGN), vinculada con indicadores sanitarios del Censo 2022 y Catálogo Georreferenciado de Efectores (Secretaría de Salud y Gestión Operativa, 2026).",
               width_in=6.4)

    # --- SECCIÓN 6: INFRAESTRUCTURA SANITARIA BÁSICA Y DETERMINANTES SOCIALES ---
    add_h1("6. Infraestructura Sanitaria Básica y Determinantes Sociales de la Salud")
    add_p("Las condiciones habitacionales y de saneamiento básico constituyen determinantes estructurales en la prevención de morbilidad infectocontagiosa infantil (diarreas, hepatitis, parasitosis). En este ámbito, el Censo 2022 ratifica que Tres de Febrero cuenta con una cobertura de infraestructura superior a la gran mayoría de los distritos del conurbano:", "Saneamiento y Calidad de Vida:")
    add_p("• Agua por red pública (cañería dentro de la vivienda): 96,1% de cobertura distrital frente al 78,4% en los 24 partidos del GBA.\n• Desagüe cloacal a red pública: 92,6% de cobertura distrital frente al 61,2% en el GBA.\n• Gas natural de red pública: 89,4% de cobertura frente al 72,1% en el GBA.\n• Conexión hogareña a Internet: 87,5% de los hogares cubiertos frente al 74,8% en el GBA.", "Indicadores de Habitabilidad y Servicios:")
    
    insert_img('grafico_cobertura_vs_infraestructura_3f.png', 
               'Gráfico 3: Comparativa de Cobertura de Salud e Infraestructura Sanitaria Básica (Tres de Febrero vs Conurbano GBA).',
               fuente="INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Hogares en viviendas particulares según acceso a servicios básicos de red de agua, cloacas, gas domiciliario e internet.")

    doc.add_page_break()

    # --- SECCIÓN 7: RED INTEGRADA DE SERVICIOS DE SALUD Y GESTIÓN MULTINIVEL ---
    add_h1("7. Red Integrada de Servicios de Salud y Gestión Multinivel")
    add_p("El municipio de Tres de Febrero cuenta con una red de servicios de salud interconnected que articula los tres niveles de atención mediante la convivencia sinérgica de efectores municipales, provinciales (Región Sanitaria VII) y nacionales, sumando además el servicio de Telemedicina municipal disponible las 24 hs.", "Modelo de Atención Integrada:")
    
    insert_img('grafico_capacidad_efectores_3f.png', 
               'Gráfico 4: Estructura Prestacional y Capacidad de la Red Sanitaria de Tres de Febrero y Región VII.',
               fuente="Relevamiento Institucional de la Secretaría de Salud y Gestión Operativa (Municipalidad de Tres de Febrero, 2026), Catálogo Federal de Establecimientos (SISA) e Informes del Ministerio de Salud PBA (Región VII).")
    
    add_h2("A. Primer Nivel y Atención Ambulatoria Especializada (Gestión Municipal)")
    add_p("• 13 Centros de Atención Primaria de la Salud (CAPS): Distribuidos en todo el territorio distrital. Horario de atención: lunes a viernes de 8 a 19 hs (en temporada de invierno de 8 a 18 hs). Brindan medicina general, clínica, pediatría, obstetricia, ginecología, psicología, psicopedagogía, trabajo social, enfermería, odontología, nutrición y dispensa de medicación.\n• Centro de Especialidades Médicas Ambulatorias de Referencia (CEMAR): Ubicado en Caseros Centro (Labardén y Labardén). Absorbe las especialidades médicas de mayor complejidad ambulatoria y se accede a ellas exclusivamente mediante derivación desde los 13 CAPS.", "CAPS y CEMAR:")
    add_p("De acuerdo a las necesidades evaluadas por los equipos de salud territorial en cada zona, en los CAPS y CEMAR se brindan espacios comunitarios (con admisión previa) tales como: 1. Ludotecas, 2. Estimulación cognitiva, 3. Grupos de habilidades socio-emocionales, 4. Talleres para mujeres y diversidades, 5. Huertas comunitarias, 6. Grupos de crianza, 7. Alfabetización, 8. Grupos de orientación a familias, 9. Caminatas saludables.", "Espacios Comunitarios y de Promoción de la Salud:")
    add_fuente_nota("Relevamiento de Programas Comunitarios, Evaluación Territorial de Admisión Previa y Directorio Operativo del Primer Nivel de Atención (Secretaría de Salud y Gestión Operativa, Municipalidad de Tres de Febrero, 2026).")

    add_p("El municipio dispone de 2 hospitales monovalentes de gestión municipal que complementan la atención primaria en especialidades de muy alta demanda: el Hospital Municipal Odontológico y el Hospital Municipal Oftalmológico 'Dr. Norberto Di Próspero', ambos con sede en Caseros Centro.", "Hospitales Monovalentes Municipales:")

    add_h2("B. Segundo y Tercer Nivel — Hospitales Provinciales (Región Sanitaria VII)")
    add_p("Dentro del territorio distrital funcionan 2 hospitales generales de gestión provincial y 1 Unidad de Pronta Atención, pertenecientes a la Región Sanitaria VII del Ministerio de Salud de la Provincia de Buenos Aires:", "Sostén de Agudos y Cuidados Críticos:")
    add_p("1. H.Z.G.A. Dr. Carlos A. Bocalandro: Av. Eva Perón (ex Ruta 8) Km 20,5 N° 9100, Loma Hermosa. Inaugurado en 1996; cuenta con aproximadamente 175 camas. Establecimiento de alta complejidad con servicios de cirugía videolaparoscópica, endoscopía, diagnóstico por imágenes, neonatología y oncoginecología. Registra del orden de 220.000 consultas externas y unos 2.000 nacimientos anuales. Es el ancla sanitaria vital del corredor norte distrital.\n\n2. H.Z.G.A. Prof. Dr. Ramón Carrillo: Hipólito Yrigoyen 1051 (y Besares), Ciudadela. Cuenta con entre 164 y 200 camas según las ampliaciones realizadas. Dispone de servicios de maternidad y neonatología (del orden de 900 a 1.200 partos anuales) y es la sede central operativa de la base del Sistema Integrado de Emergencia Sanitaria (SIES) de la Región Sanitaria VII.\n\n3. UPA 24 hs N° 16 Martín Coronado: Perón y San Lorenzo, Martín Coronado. Unidad de Pronta Atención 24 horas que estabiliza urgencias intermedias e interconecta los CAPS centrales con las guardias provinciales.", "Detalle Operativo Regional:")
    add_fuente_nota("Sistema Integrado de Información Sanitaria Argentino (SISA), Registro Federal de Establecimientos de Salud (Refes) e Informes Operativos de la Dirección de Hospitales de la Región Sanitaria VII (Ministerio de Salud PBA).")

    add_h2("C. Hospital Nacional de Referencia de Alta Complejidad")
    add_p("El Hospital Nacional 'Prof. Alejandro Posadas' constituye el establecimiento de máxima complejidad de referencia suprarregional. Ubicado en Av. Pte. Illia s/n y Av. Marconi, localidad de El Palomar (partido de Morón, limítrofe con Tres de Febrero), depende del Ministerio de Salud de la Nación bajo el régimen de Hospitales Públicos de Gestión Descentralizada. Cuenta con aproximadamente 488 camas y su área de influencia directa abarca 15 partidos de las Regiones Sanitarias V, VII y XII, siendo el centro receptor de casos críticos o derivaciones hiper-especializadas desde Tres de Febrero.", "Hospital Posadas (El Palomar):")
    add_fuente_nota("Ministerio de Salud de la Nación, Dirección Nacional de Hospitales e Indicadores Prestacionales del Hospital Nacional Prof. Alejandro Posadas.")

    # --- SECCIÓN 8: PROPUESTAS DE ARTICULACIÓN CON REGIÓN VII ---
    add_h1("8. Plan Operativo y Propuestas para la Región Sanitaria VII (2026–2027)")
    add_p("1. Plataforma de Interconsulta y Turnos Protegidos (SIREC 3F – Región VII): Formalizar un canal digital de interoperabilidad entre los 13 CAPS municipales y los consultorios externos del Hospital Bocalandro y Hospital Carrillo para derivaciones ágiles de pacientes sin cobertura formal (27,2%).\n\n2. Operativos Integrales Focalizados en el Corredor Norte: Articular rondas intersectoriales conjuntas (Municipio + Bocalandro + UPA 16) en los radios críticos de El Libertador, Churruca, Pablo Podestá, Remedios de Escalada y Loma Hermosa.\n\n3. Red Gerontológica Regional (PAMI + Hospitales + Municipio): Constituida en respuesta al 20,8% de población jubilada (75.242 personas), coordinando los talleres cognitivos y preventivos de los CAPS/CEMAR con las prestaciones interconsultoras de PAMI y los servicios geriátricos del Bocalandro y Carrillo.\n\n4. Interoperabilidad de Despacho en Emergencias (107 Municipal – Base SIES R7 Carrillo): Conectar el Centro de Monitoreo distrital con la base del SIES R7 del Hospital Carrillo de Ciudadela para optimizar tiempos de respuesta en incidentes en vía pública y emergencias cardiovasculares.", "Ejes Estratégicos Institucionales:")
    add_fuente_nota("Plan Estratégico de Salud Distrital 2026–2027 y Acuerdos Marco de Gobernanza Sanitaria Intergubernamental para la Región Sanitaria VII (Ministerio de Salud de la Provincia de Buenos Aires y Secretaría de Salud de Tres de Febrero).")
    
    # --- SECCIÓN 9: BIBLIOGRAFÍA Y FUENTES OFICIALES DE INFORMACIÓN (ENLACES DE ACCESO DIRECTO) ---
    add_h1("9. Bibliografía y Fuentes Oficiales de Información (Enlaces de Acceso Directo)")
    add_p("El presente Análisis de Situación de Salud (ASIS) del Partido de Tres de Febrero ha sido elaborado bajo estándares metodológicos rigurosos, consolidando bases de datos primarias y secundarias de organismos estadísticos y sanitarios de jerarquía nacional, provincial y municipal. A continuación, se detallan las referencias bibliográficas oficiales y los enlaces de acceso web directo para la consulta pública y auditoría de la información:", "Fuentes y Enlaces Institucionales:")
    
    bibliografia_items = [
        ("INDEC — Censo Nacional de Población, Hogares y Viviendas 2022:",
         "Instituto Nacional de Estadística y Censos (INDEC). Resultados definitivos del Censo 2022: Indicadores demográficos por sexo, grupos de edad, fecundidad y características estructurales de las viviendas particulares para el Partido de Tres de Febrero (Código 06840).\nEnlace de acceso directo: https://www.indec.gob.ar/indec/web/Nivel4-Tema-2-41-165"),
        ("INDEC — Geoportal y Cartografía Georreferenciada por Radios Censales:",
         "Base cartográfica poligonal oficial y cobertura sanitaria por radios censales e intervalos espaciales. Sistema de Consulta Georreferenciada INDEC.\nEnlace de acceso directo: https://geoah.indec.gob.ar/ y https://www.indec.gob.ar/indec/web/Institucional-Indec-Censo2022-MarcoEstadistico"),
        ("SISA / REFES — Registro Federal de Establecimientos de Salud:",
         "Ministerio de Salud de la Nación. Sistema Integrado de Información Sanitaria Argentino (SISA) y Catálogo Nacional de Establecimientos de Salud (REFES). Georreferenciación, tipología prestacional y dependencia de hospitales provinciales, nacionales y los 13 CAPS municipales de Tres de Febrero.\nEnlace de acceso directo: https://sisa.msal.gov.ar/sisa/#sika"),
        ("Ministerio de Salud PBA — Región Sanitaria VII:",
         "Dirección Provincial de Hospitales y Coordinación de la Región Sanitaria VII (Morón, Tres de Febrero, Hurlingham, Ituzaingó, Merlo, Moreno, Gral. Las Heras, Gral. Rodríguez, Luján y Marcos Paz). Informes epidemiológicos y red de emergencias SIES.\nEnlace de acceso directo: https://www.ms.gba.gov.ar/sitios/regiones-sanitarias/ y https://www.ms.gba.gov.ar/sitios/hospitales/"),
        ("DEIS — Dirección de Estadísticas e Información en Salud:",
         "Ministerio de Salud de la Nación. Estadísticas vitales, mortalidad infantil, materna y por causas en el Conurbano Bonaerense y Región Sanitaria VII.\nEnlace de acceso directo: https://deis.msal.gov.ar/"),
        ("Secretaría de Salud y Gestión Operativa — Municipalidad de Tres de Febrero:",
         "Portal de Gobierno Municipal, red de Centros de Atención Primaria de la Salud (CAPS 1 al 13), CEMAR y Hospital Odontológico y Oftalmológico Di Próspero. Programas de salud territorial y turnero municipal.\nEnlace de acceso directo: https://www.tresdefebrero.gov.ar/salud/"),
        ("PAMI — Instituto Nacional de Servicios Sociales para Jubilados y Pensionados:",
         "Estadísticas de cobertura prestacional gerontológica, cápitas asignadas por clínica/hospital e interconsultas especializadas para adultos mayores (≥65 años).\nEnlace de acceso directo: https://www.pami.org.ar/ y https://estadisticas.pami.org.ar/"),
        ("OpenStreetMap Nominatim — Geocodificación Geoespacial Verificada:",
         "Base de datos cartográfica global para verificación de coordenadas exactas y parcelas catastrales de establecimientos de salud en el Corredor Norte, Centro y Sur de Tres de Febrero.\nEnlace de acceso directo: https://nominatim.openstreetmap.org/")
    ]
    
    for titulo_b, desc_b in list(bibliografia_items):
        p_bib = doc.add_paragraph()
        p_bib.paragraph_format.space_after = Pt(8)
        p_bib.paragraph_format.line_spacing = 1.15
        run_t = p_bib.add_run(f"• {titulo_b} ")
        run_t.bold = True
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(22, 60, 104) # Azul marino
        
        run_d = p_bib.add_run(desc_b)
        run_d.font.name = 'Calibri'
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = RGBColor(51, 51, 51)
    
    out_docx_completo = os.path.join(base_dir, 'informe_situacion_salud_tres_de_febrero_COMPLETO.docx')
    out_docx_actualizado = os.path.join(base_dir, 'informe_situacion_salud_tres_de_febrero_ACTUALIZADO.docx')
    out_docx_salidas = os.path.join(base_dir, 'salidas', 'informe_situacion_salud_tres_de_febrero_COMPLETO.docx')
    
    os.makedirs(os.path.join(base_dir, 'salidas'), exist_ok=True)
    try:
        doc.save(out_docx_salidas)
        print(f" -> Guardado en salidas: {out_docx_salidas}")
    except Exception as e:
        print(f" [Nota] No se pudo guardar en salidas: {e}")

    try:
        doc.save(out_docx_completo)
        print(f"[¡ÉXITO DOCX UNIFICADO CON REFERENCIAS ACADÉMICAS EN TODAS LAS TABLAS Y GRÁFICOS!] Guardado en:\n -> {out_docx_completo}")
    except PermissionError:
        print(f"[¡NOTA IMPORTANTE!] {os.path.basename(out_docx_completo)} estaba abierto en tu Microsoft Word.")

    try:
        doc.save(out_docx_actualizado)
        print(f" -> Guardado en: {out_docx_actualizado}")
    except Exception as e:
        print(f" [Nota] No se pudo guardar ACTUALIZADO: {e}")

    try:
        doc.save(out_docx)
        print(f" -> También actualizado: {out_docx}")
    except PermissionError:
        pass

if __name__ == '__main__':
    generar_word_asis()
