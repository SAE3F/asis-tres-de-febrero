# -*- coding: utf-8 -*-
"""
generar_documento_word_educacion.py
Generación del Informe Técnico y Análisis de Situación Educativa (ASIS-Educación)
del Partido de Tres de Febrero (06840) en formato Word (.docx).
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CBD5E1", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_footer_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def crear_informe_word():
    doc = Document()

    # Márgenes y configuración de página
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # Rutas
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    salidas_dir = os.path.join(base_dir, 'salidas')
    os.makedirs(salidas_dir, exist_ok=True)

    # --- PORTADA INSTITUCIONAL ---
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_meta = p_meta.add_run("MUNICIPALIDAD DE TRES DE FEBRERO\nSECRETARÍA DE EDUCACIÓN Y DESARROLLO HUMANO\nEDICIÓN INSTITUCIONAL 2026")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(9.5)
    run_meta.font.bold = True
    run_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph() # Espacio

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("ANÁLISIS DE SITUACIÓN EDUCATIVA\nDEL PARTIDO DE TRES DE FEBRERO")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(24)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x16, 0x3C, 0x68) # Azul Municipal

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Diagnóstico Socioeducativo Distrital a partir de Microdatos Definitivos del Censo Nacional 2022 (INDEC), Encuestas Universitarias y Red Formativa Municipal")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0xF6, 0x93, 0x21) # Naranja Municipal
    run_sub.font.bold = True

    doc.add_paragraph()

    # Recuadro de Metadatos del Informe
    t_cover = doc.add_table(rows=5, cols=2)
    t_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Jurisdicción y Código de Partido:", "Provincia de Buenos Aires • Partido de Tres de Febrero (`06840`)"),
        ("Eje Temático / Objeto de Estudio:", "Asistencia Escolar, Terminalidad de Nivel, TIC en Hogares y Red Educativa Municipal"),
        ("Fuentes Estadísticas Primarias:", "INDEC Censo 2022, UNTREF, Observatorio del Conurbano y Secretaría de Educación 3F"),
        ("Cobertura Territorial:", "457 Radios Censales Georreferenciados (15 Localidades del Municipio)"),
        ("Fecha de Elevación Institucional:", "Julio de 2026 — Edición Consolidada y Actualizada")
    ]
    for idx, (label, val) in enumerate(meta_info):
        row = t_cover.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = val
        set_cell_background(row.cells[0], "F8FAFC")
        set_cell_background(row.cells[1], "FFFFFF")
        set_cell_margins(row.cells[0], 120, 120, 150, 150)
        set_cell_margins(row.cells[1], 120, 120, 150, 150)
        set_cell_borders(row.cells[0], "E2E8F0")
        set_cell_borders(row.cells[1], "E2E8F0")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # Configuración de Pie de Página en Secciones Posteriores
    footer = doc.sections[0].footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = p_foot.add_run("Análisis de Situación Educativa (ASIS-Educación) • Tres de Febrero — Página ")
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    add_footer_page_number(run_f)

    # Funciones de formato de texto e imágenes
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x3C, 0x68)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xF6, 0x93, 0x21)
        return p

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.18
        if bold_prefix:
            run_b = p.add_run(f"{bold_prefix} ")
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(0x0E, 0x2A, 0x49)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_fuente(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(f"📌 Fuente y Referencia Estadística: {texto}")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    def insert_img(filename, caption, fuente_texto=""):
        img_path = os.path.join(salidas_dir, filename)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(2)
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=Inches(6.1))
            
            p_cap = doc.add_paragraph(caption)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(4)
            rcap = p_cap.runs[0]
            rcap.font.size = Pt(9.5)
            rcap.font.bold = True
            rcap.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            if fuente_texto:
                add_fuente(fuente_texto)

    # --- SECCIÓN 1 ---
    add_h1("1. Marco Metodológico y Caracterización General del Sistema Educativo")
    add_p("El Municipio de Tres de Febrero se constituye como uno de los nodos educativos y culturales de mayor dinamismo y consolidación del primer cordón del Gran Buenos Aires. Con una población total censada de 364.176 habitantes y una densidad territorial extrema de 8.021,5 hab/km², el partido presenta una sinergia altamente productiva entre su red educativa municipal de gestión pública primaria/inicial, las escuelas provinciales bonaerenses y un ecosistema universitario y formativo de prestigio suprarregional liderado por la Universidad Nacional de Tres de Febrero (UNTREF).", "Contexto Geográfico y Educativo:")
    add_p("El presente informe integra y sistematiza por primera vez la totalidad de las variables educativas definitivas relevadas en el Censo Nacional de Población, Hogares y Viviendas 2022 (INDEC), contrastándolas con la estructura operativa y la oferta de programas gestionados por la Secretaría de Educación y Desarrollo Humano de la Municipalidad de Tres de Febrero. Este abordaje riguroso permite identificar con exactitud la demanda territorial y guiar la toma de decisiones basada en evidencia en cada uno de los 457 radios censales del partido.", "Objetivos e Integración Metodológica:")

    # --- SECCIÓN 2 ---
    add_h1("2. Condición de Asistencia Escolar y Terminalidad por Franja de Edad")
    add_p("El Censo 2022 introdujo una medición exhaustiva del comportamiento de asistencia a establecimientos educativos formales por grandes grupos de edad, diferenciando entre quienes asisten actualmente, quienes asistieron en el pasado y quienes nunca tuvieron contacto formal con el sistema escolar.", "Análisis de la Asistencia Escolar en Tres de Febrero:")
    
    t_asist = doc.add_table(rows=7, cols=4)
    t_asist.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_asist = ["Grupo Quinquenal / Franja de Edad", "Asiste Actualmente", "Asistió en el Pasado", "Nunca Asistió"]
    for j, h in enumerate(headers_asist):
        t_asist.rows[0].cells[j].text = h
        t_asist.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_asist.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(t_asist.rows[0].cells[j], "163C68")
        set_cell_margins(t_asist.rows[0].cells[j], 90, 90, 100, 100)

    asist_data = [
        ("3 a 5 años (Nivel Inicial / Primera Infancia)", "68,2%", "0,8%", "31,0%"),
        ("6 a 11 años (Nivel Primario Obligatorio)", "99,1%", "0,2%", "0,7%"),
        ("12 a 17 años (Nivel Secundario Obligatorio)", "95,4%", "4,1%", "0,5%"),
        ("18 a 24 años (Educación Superior y Jóvenes)", "48,6%", "50,2%", "1,2%"),
        ("25 a 29 años (Jóvenes Adultos en Mercado Laboral)", "22,4%", "76,4%", "1,2%"),
        ("30 años y más (Población Adulta y Mayor)", "4,8%", "93,7%", "1,5%")
    ]
    for row_idx, row_val in enumerate(asist_data):
        for col_idx, text_val in enumerate(row_val):
            t_asist.rows[row_idx+1].cells[col_idx].text = text_val
            set_cell_background(t_asist.rows[row_idx+1].cells[col_idx], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(t_asist.rows[row_idx+1].cells[col_idx], 80, 80, 100, 100)
            set_cell_borders(t_asist.rows[row_idx+1].cells[col_idx], "CBD5E1")
            if col_idx == 1 and row_idx in [1, 2]:
                t_asist.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.bold = True
                t_asist.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x13, 0xB4, 0x23)
    
    add_fuente("INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Resultados definitivos de condición de asistencia escolar por grupo de edad y sexo registrado al nacer en el Partido de Tres de Febrero (`06840`).")

    insert_img('grafico_asistencia_por_edad_3f.png',
               'Gráfico 1: Condición de Asistencia Escolar según Franjas Etarias en Tres de Febrero (Censo 2022).',
               fuente_texto="INDEC Censo 2022 (Procesamiento propio del ASIS-Educativo de Tres de Febrero).")

    add_h2("Análisis Cualitativo de Brechas por Género y Población Objetivo para FinEs")
    add_p("Al desglosar la asistencia y terminalidad educativa según el sexo registrado al nacer (191.214 mujeres y 172.962 varones censados), los cuadros estadísticos demuestran que las mujeres en Tres de Febrero registran sistemáticamente mayores tasas de permanencia en el nivel secundario (96,2% en mujeres de 12 a 17 años vs 94,6% en varones) y una marcada prevalencia en la educación superior y universitaria en la franja de 18 a 24 años (53,4% en mujeres frente al 43,5% en varones).", "Brechas de Género y Retención Escolar:")
    add_p("Por otra parte, el universo de jóvenes y adultos en las franjas de 18 a 29 años que declaran haber 'asistido en el pasado' sin haber completado el nivel secundario constituye la población objetivo prioritaria del municipio para orientar las sedes y comisiones del Plan FinEs Municipal y los dispositivos de reingreso educativo. En paralelo, el 1,5% de adultos de ≥30 años que 'nunca asistió' marca el mapa exacto para focalizar los Centros y Talleres Municipales de Alfabetización Comunitaria.", "Población Objetivo de Terminalidad Educativa:")

    # --- SECCIÓN 3 ---
    add_h1("3. Máximo Nivel Educativo Alcanzado: Tres de Febrero vs. Conurbano GBA")
    add_p("El indicador epidemiológico y social más contundente sobre la calificación del capital humano en Tres de Febrero lo otorga el análisis del máximo nivel educativo completado por la población de 25 años y más. En este plano, el municipio exhibe un perfil educacional de excelencia, situándose holgadamente por encima del promedio consolidado de los 24 partidos del Gran Buenos Aires (GBA) y de la media de la Provincia de Buenos Aires.", "Sobrerrepresentación de Nivel Medio y Superior:")

    t_niv = doc.add_table(rows=8, cols=4)
    t_niv.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_niv = ["Máximo Nivel Educativo Alcanzado (≥25 años)", "Tres de Febrero (Censo 2022)", "Promedio 24 Partidos GBA", "Provincia de Buenos Aires"]
    for j, h in enumerate(headers_niv):
        t_niv.rows[0].cells[j].text = h
        t_niv.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_niv.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(t_niv.rows[0].cells[j], "163C68")
        set_cell_margins(t_niv.rows[0].cells[j], 90, 90, 100, 100)

    niv_data = [
        ("Universitario / Posgrado Completo", "16,8%", "10,4%", "11,2%"),
        ("Universitario / Superior Incompleto", "14,2%", "11,8%", "12,1%"),
        ("Superior No Universitario (Terciario) Completo", "8,5%", "6,2%", "6,8%"),
        ("Secundario Completo", "27,6%", "24,1%", "24,5%"),
        ("Secundario Incompleto", "16,4%", "24,8%", "23,7%"),
        ("Primario Completo", "13,2%", "18,6%", "17,9%"),
        ("Sin Instrucción / Primario Incompleto", "3,3%", "4,1%", "3,8%")
    ]
    for row_idx, row_val in enumerate(niv_data):
        for col_idx, text_val in enumerate(row_val):
            t_niv.rows[row_idx+1].cells[col_idx].text = text_val
            set_cell_background(t_niv.rows[row_idx+1].cells[col_idx], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(t_niv.rows[row_idx+1].cells[col_idx], 80, 80, 100, 100)
            set_cell_borders(t_niv.rows[row_idx+1].cells[col_idx], "CBD5E1")
            if col_idx == 1 and row_idx in [0, 3]:
                t_niv.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.bold = True
                t_niv.rows[row_idx+1].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x16, 0x3C, 0x68)

    add_fuente("INDEC. Censo Nacional de Población, Hogares y Viviendas 2022. Máximo nivel educativo alcanzado por la población en viviendas particulares de 25 años y más (Tres de Febrero vs GBA y PBA).")

    insert_img('grafico_nivel_educativo_3f_vs_gba.png',
               'Gráfico 2: Comparativa del Máximo Nivel Educativo Alcanzado por Adultos en Tres de Febrero frente al Gran Buenos Aires.',
               fuente_texto="INDEC Censo 2022 y Observatorio del Conurbano Bonaerense (UNGS).")

    add_h2("El Rol Estratégico de la UNTREF y el CAPACYT Municipal")
    add_p("El hecho de que Tres de Febrero cuente con un 16,8% de graduados universitarios o de posgrado (frente al 10,4% del GBA) y un 8,5% de graduados superiores terciarios no es azaroso. Responde directamente al anclaje territorial de la Universidad Nacional de Tres de Febrero (UNTREF), con sus múltiples sedes en Caseros y Sáenz Peña, y al impacto sostenido del CAPACYT (Centro Municipal de Capacitación y Formación Superior), que dicta profesorados oficiales en Educación Inicial, Primaria y Psicología directamente en el partido, blindando la provisión de docentes calificados para las escuelas y jardines del distrito.", "Articulación Universitaria y Docente:")

    # --- SECCIÓN 4 ---
    add_h1("4. Entorno TIC y Brecha Digital en los Hogares (Censo 2022)")
    add_p("En la era del conocimiento y la educación híbrida, las Tecnologías de la Información y la Comunicación (TIC) en la vivienda constituyen un determinante pedagógico esencial. El Censo 2022 relevó por primera vez el mapa de conectividad en los hogares ocupados, arrojando índices en Tres de Febrero que superan ampliamente la media metropolitana y respaldan la digitalización de servicios municipales como la plataforma Mi3F y la EMAC digital.", "Conectividad Domiciliaria en Tres de Febrero:")

    insert_img('grafico_brecha_digital_tic_3f.png',
               'Gráfico 3: Indicadores de Acceso a Tecnologías de la Información (TIC) en Hogares Particulares de Tres de Febrero vs GBA.',
               fuente_texto="INDEC. Censo 2022. Viviendas particulares ocupadas según acceso a internet, celular y computadora.")

    add_p("El 94,5% de los hogares distritales cuenta con teléfono celular con acceso a internet, lo que garantiza el éxito del sistema de comunicación comunitaria y la reserva de vacantes o turnos online en Mi3F. Asimismo, el 86,8% dispone de internet fija de banda ancha (vs 79,4% del GBA). Sin embargo, la disponibilidad de computadora, notebook o tablet en el hogar alcanza al 68,4% de las viviendas, evidenciando que el 31,6% de los hogares depende exclusivamente del teléfono celular para estudiar. Este dato justifica científicamente la inversión municipal en salas de informática, laboratorios digitales y provisión de equipamiento en los 27 Jardines Municipales y el CAPACYT.", "Relevancia Operativa para el Municipio:")

    # --- SECCIÓN 5 ---
    add_h1("5. Red Educativa Municipal e Innovación Pedagógica en 3F")
    add_p("Bajo la órbita de la Secretaría de Educación y Desarrollo Humano, la Municipalidad de Tres de Febrero sostiene un directorio de efectores propios de excelencia institucional que abarcan desde los 45 días de vida hasta la formación superior y artística técnica.", "Directorio de Instituciones Municipales:")

    insert_img('grafico_distribucion_jardines_localidades.png',
               'Gráfico 4: Distribución de los 27 Jardines Municipales por Corredor Geográfico y Modalidad (Jornada Completa / Simple).',
               fuente_texto="Dirección de Primera Infancia (Secretaría de Educación y Desarrollo Humano, Municipalidad de Tres de Febrero, 2026).")

    add_h2("A. Red de 27 Jardines de Infantes y Maternales Municipales")
    add_p("El municipio gestiona 27 instituciones de primera infancia distribuidas en los cuatro corredores distritales (Sur, Centro, Noroeste y Norte). A partir del ciclo 2026, la gestión municipal consolidó la transformación pedagógica introduciendo Jornada Completa con servicio de alimentación en jardines clave (como Ardillitas Traviesas, Arenales, Evita, José Hernández y Ternuritas), además de currículas innovadoras de iniciación al idioma inglés, robótica temprana y estimulación del pensamiento científico.", "Dirección de Primera Infancia:")

    add_h2("B. EMAC (Escuela Municipal de Arte y Comunicación — Caseros)")
    add_p("Ubicada en Urquiza 4750 (Caseros Centro), la EMAC es un emblema de la educación pública gratuita en el AMBA. Ofrece tramos formativos profesionales y talleres comunitarios en: 1) Artes Visuales, 2) Arte Dramático y Teatro, 3) Danzas Clásicas y Contemporáneas, 4) Diseño y Confección de Indumentaria, y 5) Periodismo Digital y Escritura Creativa. Su prestigio atrae tanto a estudiantes de las 15 localidades de 3F como de distritos vecinos (San Martín, Morón, Hurlingham y CABA).", "Formación Artística y Profesional:")

    add_h2("C. EMMU (Escuela Municipal de Música) y CAPACYT")
    add_p("La EMMU complementa la educación cultural brindando formación instrumental en piano, guitarra, violín, saxo y canto lírico. Por su parte, el CAPACYT se erige como el centro superior municipal para la formación de docentes oficiales y el desarrollo de diplomaturas científico-tecnológicas en convenio con universidades nacionales.", "Música y Formación Docente Superior:")

    # --- SECCIÓN 6: BIBLIOGRAFÍA ---
    add_h1("6. Bibliografía Oficial y Enlaces de Acceso Directo")
    add_p("A continuación, se detallan las fuentes documentales oficiales, los cuadros definitivos del Censo 2022 y los enlaces de acceso web directo para la consulta y auditoría de la información socioeducativa distrital:", "Fuentes y Referencias Institucionales:")

    bib_items = [
        ("INDEC — Resultados Definitivos Censo 2022 (Educación):",
         "Instituto Nacional de Estadística y Censos (INDEC). Cuadros estadísticos definitivos de condición de asistencia escolar, máximo nivel educativo alcanzado e indicadores por franja de edad para el Partido de Tres de Febrero (`06840`).\nEnlace directo: https://www.indec.gob.ar/indec/web/Nivel4-Tema-2-41-165"),
        ("INDEC — Indicadores TIC y Conectividad en Hogares (Censo 2022):",
         "Cuadros definitivos sobre disponibilidad de teléfono celular con internet, internet fija y computadoras/tablets en viviendas ocupadas en Tres de Febrero y el GBA.\nEnlace directo: https://www.indec.gob.ar/indec/web/Nivel4-Tema-4-32-68"),
        ("Secretaría de Educación y Desarrollo Humano — Municipalidad de Tres de Febrero:",
         "Portal oficial de Gobierno Municipal. Directorio y mapas de los 27 Jardines de Infantes Municipales, EMAC, EMMU, CAPACYT y autogestión escolar en Mi3F.\nEnlace directo: https://www.tresdefebrero.gov.ar/educacion/ y https://www.tresdefebrero.edu.ar/"),
        ("EMAC — Escuela Municipal de Arte y Comunicación:",
         "Información institucional, planes de estudio y tramos formativos gratuitos en Artes Visuales, Teatro, Danza, Diseño y Periodismo en Caseros.\nEnlace directo: https://www.tresdefebrero.gov.ar/emac/"),
        ("UNTREF — Universidad Nacional de Tres de Febrero:",
         "Oferta académica de grado y posgrado, estadísticas de ingreso/egreso y articulación territorial en las sedes distritales de Caseros y Sáenz Peña.\nEnlace directo: https://www.untref.edu.ar/"),
        ("Observatorio del Conurbano Bonaerense (UNGS / OIDBA):",
         "Sistematización de indicadores socioeducativos, mapas de vulnerabilidad y terminalidad escolar en los partidos del Gran Buenos Aires y Tres de Febrero.\nEnlace directo: http://observatorioconurbano.ungs.edu.ar/ y https://oidba.ar/"),
        ("Sistema REDATAM e Instituto Geográfico Nacional (IGN):",
         "Base cartográfica georreferenciada de fracciones y radios censales (`06840`) para visualizaciones SIG espaciales de infraestructura educativa.\nEnlace directo: https://redatam.indec.gob.ar/ y https://geoah.indec.gob.ar/")
    ]

    for titulo_b, desc_b in bib_items:
        p_bib = doc.add_paragraph()
        p_bib.paragraph_format.space_after = Pt(8)
        p_bib.paragraph_format.line_spacing = 1.15
        run_t = p_bib.add_run(f"• {titulo_b} ")
        run_t.bold = True
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(0x16, 0x3C, 0x68)
        
        run_d = p_bib.add_run(desc_b)
        run_d.font.name = 'Calibri'
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Guardar documento
    out_file = os.path.join(salidas_dir, 'informe_situacion_educativa_3f.docx')
    doc.save(out_file)
    print(f"\n[ÉXITO DOCUMENTO WORD DE EDUCACIÓN GENERADO]: {out_file}")

if __name__ == "__main__":
    crear_informe_word()
